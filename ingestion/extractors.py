"""Document text extraction utilities for various file formats."""

import io
from pathlib import Path
from typing import Dict, List, Optional
import pdfplumber
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pytesseract
from PIL import Image
from loguru import logger


class DocumentExtractor:
    """Base class for document text extraction."""

    @staticmethod
    def extract(file_path: Path) -> Dict[str, any]:
        """Extract text from a document.

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary containing extracted text and metadata
        """
        raise NotImplementedError


class PDFExtractor(DocumentExtractor):
    """Extract text from PDF documents."""

    @staticmethod
    def extract(file_path: Path, use_ocr: bool = False) -> Dict[str, any]:
        """Extract text from PDF file.

        Args:
            file_path: Path to PDF file
            use_ocr: Whether to use OCR for scanned pages

        Returns:
            Dictionary with extracted text, page count, and metadata
        """
        try:
            text_pages = []
            metadata = {}

            # Try pdfplumber first (better for text-based PDFs)
            with pdfplumber.open(file_path) as pdf:
                metadata = {
                    "pages": len(pdf.pages),
                    "producer": pdf.metadata.get("Producer", ""),
                    "creator": pdf.metadata.get("Creator", ""),
                }

                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()

                    # If no text found and OCR is enabled, try OCR
                    if (not page_text or page_text.strip() == "") and use_ocr:
                        page_text = PDFExtractor._ocr_page(file_path, i)

                    text_pages.append({
                        "page_number": i + 1,
                        "text": page_text or "",
                    })

            full_text = "\n\n".join([p["text"] for p in text_pages])

            return {
                "text": full_text,
                "pages": text_pages,
                "metadata": metadata,
                "success": True,
            }

        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {e}")
            return {
                "text": "",
                "pages": [],
                "metadata": {},
                "success": False,
                "error": str(e),
            }

    @staticmethod
    def _ocr_page(file_path: Path, page_num: int) -> str:
        """Perform OCR on a specific PDF page.

        Args:
            file_path: Path to PDF file
            page_num: Page number (0-indexed)

        Returns:
            Extracted text from OCR
        """
        try:
            # Use PyMuPDF to render page as image
            doc = fitz.open(file_path)
            page = doc[page_num]

            # Render at 2x resolution for better OCR
            mat = fitz.Matrix(2, 2)
            pix = page.get_pixmap(matrix=mat)

            # Convert to PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Perform OCR
            text = pytesseract.image_to_string(img)

            doc.close()
            return text

        except Exception as e:
            logger.error(f"Error performing OCR on page {page_num}: {e}")
            return ""


class DOCXExtractor(DocumentExtractor):
    """Extract text from DOCX documents."""

    @staticmethod
    def extract(file_path: Path) -> Dict[str, any]:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            doc = DocxDocument(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name if para.style else None,
                    })

            # Extract text from tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            # Combine all text
            full_text = "\n\n".join([p["text"] for p in paragraphs])

            # Add table text
            for table in tables:
                for row in table:
                    full_text += "\n" + " | ".join(row)

            metadata = {
                "paragraphs": len(paragraphs),
                "tables": len(tables),
            }

            # Try to get core properties
            try:
                core_props = doc.core_properties
                metadata.update({
                    "author": core_props.author or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                    "title": core_props.title or "",
                })
            except:
                pass

            return {
                "text": full_text,
                "paragraphs": paragraphs,
                "tables": tables,
                "metadata": metadata,
                "success": True,
            }

        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {e}")
            return {
                "text": "",
                "paragraphs": [],
                "tables": [],
                "metadata": {},
                "success": False,
                "error": str(e),
            }


def extract_document(file_path: Path, use_ocr: bool = False) -> Dict[str, any]:
    """Extract text from a document based on file extension.

    Args:
        file_path: Path to document
        use_ocr: Whether to use OCR for scanned documents

    Returns:
        Extracted text and metadata
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return PDFExtractor.extract(file_path, use_ocr=use_ocr)
    elif suffix in [".docx", ".doc"]:
        return DOCXExtractor.extract(file_path)
    else:
        logger.warning(f"Unsupported file format: {suffix}")
        return {
            "text": "",
            "metadata": {},
            "success": False,
            "error": f"Unsupported file format: {suffix}",
        }
